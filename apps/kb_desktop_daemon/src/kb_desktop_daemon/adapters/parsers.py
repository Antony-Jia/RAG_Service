from __future__ import annotations

import io
from pathlib import Path

from docx import Document as DocxDocument
from kb_core.models import BlobRef, ParsedDocument, ParseOptions
from pypdf import PdfReader


class TextParser:
    def can_parse(self, mime: str, ext: str) -> bool:
        return mime.startswith("text/") or ext in {".txt", ".md"}

    def parse(self, blob: BlobRef, opts: ParseOptions) -> ParsedDocument:
        content = Path(blob.path).read_text(encoding="utf-8", errors="ignore")
        return ParsedDocument(text=content, metadata={"parser": "text"})


class PdfParser:
    def can_parse(self, mime: str, ext: str) -> bool:
        return mime == "application/pdf" or ext == ".pdf"

    def parse(self, blob: BlobRef, opts: ParseOptions) -> ParsedDocument:
        reader = PdfReader(blob.path)
        lines: list[str] = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                lines.append(text)
            lines.append(f"\n[page={page_index}]\n")
        return ParsedDocument(text="\n".join(lines).strip(), metadata={"parser": "pdf"})


class DocxParser:
    def can_parse(self, mime: str, ext: str) -> bool:
        return (
            mime
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or ext == ".docx"
        )

    def parse(self, blob: BlobRef, opts: ParseOptions) -> ParsedDocument:
        file_bytes = Path(blob.path).read_bytes()
        doc = DocxDocument(io.BytesIO(file_bytes))
        text = "\n".join(para.text for para in doc.paragraphs if para.text)
        return ParsedDocument(text=text, metadata={"parser": "docx"})


def default_parsers() -> list[object]:
    return [TextParser(), PdfParser(), DocxParser()]
